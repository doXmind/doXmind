"""Export service for converting Markdown to PDF and DOCX.

This module provides multi-format export with a unified HTML parsing layer
to avoid code duplication between formats.
"""

from __future__ import annotations

import contextlib
import io
import logging
import math
import os
import re
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx.document import Document
    from fpdf import FPDF

logger = logging.getLogger(__name__)

# ============================================================================
# Document Node Model (Intermediate Representation)
# ============================================================================


class NodeType(Enum):
    """Types of document nodes."""

    TEXT = "text"
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    CODE_BLOCK = "code_block"
    BLOCKQUOTE = "blockquote"
    LIST = "list"
    LIST_ITEM = "list_item"
    TABLE = "table"
    TABLE_ROW = "table_row"
    TABLE_CELL = "table_cell"
    HORIZONTAL_RULE = "hr"
    MERMAID_CHART = "mermaid_chart"
    INLINE_BOLD = "bold"
    INLINE_ITALIC = "italic"
    INLINE_CODE = "code"
    INLINE_LINK = "link"


@dataclass
class DocumentNode:
    """Represents a node in the document tree."""

    node_type: NodeType
    content: str = ""
    children: list[DocumentNode] = field(default_factory=list)

    # Type-specific attributes
    level: int = 1  # For headings (1-6)
    ordered: bool = False  # For lists
    is_header: bool = False  # For table cells
    url: str = ""  # For links


@dataclass
class ExportMetadata:
    """Metadata for rendering a title page in exported documents."""

    title: str = ""
    icon: str | None = None
    author: str | None = None
    cover_image_url: str | None = None
    cover_position: float = 0.5
    created_at: datetime | None = None
    updated_at: datetime | None = None


# ============================================================================
# HTML Parser (converts HTML to DocumentNode tree)
# ============================================================================


class HTMLToDocumentParser:
    """Parses HTML into a DocumentNode tree."""

    def parse(self, html: str) -> list[DocumentNode]:
        """Parse HTML string into document nodes."""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        return self._parse_children(soup)

    def _parse_children(self, element) -> list[DocumentNode]:
        """Parse children of an element."""
        nodes = []
        for child in element.children:
            node = self._parse_element(child)
            if node:
                nodes.append(node)
        return nodes

    def _parse_element(self, element) -> DocumentNode | None:
        """Parse a single element into a DocumentNode."""
        if element.name is None:
            # Text node
            text = str(element).strip()
            if text:
                return DocumentNode(NodeType.TEXT, content=text)
            return None

        name = element.name

        # Headings
        if name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(name[1])
            return DocumentNode(NodeType.HEADING, content=element.get_text().strip(), level=level)

        # Paragraph
        if name == "p":
            return DocumentNode(NodeType.PARAGRAPH, children=self._parse_inline_elements(element))

        # Code block
        if name == "pre":
            return DocumentNode(NodeType.CODE_BLOCK, content=element.get_text())

        # Blockquote
        if name == "blockquote":
            return DocumentNode(NodeType.BLOCKQUOTE, content=element.get_text().strip())

        # Lists
        if name in ["ul", "ol"]:
            return DocumentNode(
                NodeType.LIST, ordered=(name == "ol"), children=self._parse_list_items(element)
            )

        # Table
        if name == "table":
            return DocumentNode(NodeType.TABLE, children=self._parse_table_rows(element))

        # Horizontal rule
        if name == "hr":
            return DocumentNode(NodeType.HORIZONTAL_RULE)

        # Mermaid chart
        if name == "div" and element.get("data-type") == "mermaid-chart":
            code = element.get("data-code", "")
            return DocumentNode(NodeType.MERMAID_CHART, content=code)

        # Div - recurse
        if name == "div":
            children = self._parse_children(element)
            if len(children) == 1:
                return children[0]
            elif children:
                # Wrap in paragraph if multiple children
                return DocumentNode(NodeType.PARAGRAPH, children=children)

        return None

    def _parse_inline_elements(self, element) -> list[DocumentNode]:
        """Parse inline elements (bold, italic, code, links)."""
        nodes = []
        for child in element.children:
            if child.name is None:
                text = str(child)
                if text:
                    nodes.append(DocumentNode(NodeType.TEXT, content=text))
            elif child.name in ["strong", "b"]:
                nodes.append(DocumentNode(NodeType.INLINE_BOLD, content=child.get_text()))
            elif child.name in ["em", "i"]:
                nodes.append(DocumentNode(NodeType.INLINE_ITALIC, content=child.get_text()))
            elif child.name == "code":
                nodes.append(DocumentNode(NodeType.INLINE_CODE, content=child.get_text()))
            elif child.name == "a":
                nodes.append(
                    DocumentNode(
                        NodeType.INLINE_LINK, content=child.get_text(), url=child.get("href", "")
                    )
                )
            elif child.name == "br":
                nodes.append(DocumentNode(NodeType.TEXT, content="\n"))
            else:
                # Recurse for other elements
                nodes.extend(self._parse_inline_elements(child))
        return nodes

    def _parse_list_items(self, element) -> list[DocumentNode]:
        """Parse list items."""
        items = []
        for li in element.find_all("li", recursive=False):
            # Get direct text content
            text_parts = []
            nested_lists = []

            for child in li.children:
                if child.name is None:
                    text_parts.append(str(child))
                elif child.name in ["ul", "ol"]:
                    nested_lists.append(self._parse_element(child))
                elif child.name not in ["ul", "ol"]:
                    text_parts.append(child.get_text())

            item = DocumentNode(
                NodeType.LIST_ITEM,
                content="".join(text_parts).strip(),
                children=[n for n in nested_lists if n],
            )
            items.append(item)
        return items

    def _parse_table_rows(self, element) -> list[DocumentNode]:
        """Parse table rows."""
        rows = []
        for tr in element.find_all("tr"):
            cells = []
            for cell in tr.find_all(["th", "td"]):
                cells.append(
                    DocumentNode(
                        NodeType.TABLE_CELL,
                        content=cell.get_text().strip(),
                        is_header=(cell.name == "th"),
                    )
                )
            if cells:
                rows.append(DocumentNode(NodeType.TABLE_ROW, children=cells))
        return rows


# ============================================================================
# PDF Renderer
# ============================================================================


class PDFRenderer:
    """Renders DocumentNodes to PDF."""

    def __init__(self):
        self._unicode_font = None
        self._unicode_font_path = None
        self._find_unicode_font()

    def _find_unicode_font(self):
        """Find a suitable Unicode font on the system."""
        # First try to find fonts by scanning directories (more reliable for Linux/Docker)
        font_path = self._scan_for_cjk_font()
        if font_path:
            self._unicode_font_path = font_path
            self._unicode_font = "CJK Font"
            return

        # Fallback to specific file names
        font_candidates = [
            # Linux (fonts-noto-cjk package) - various possible names
            ("NotoSansCJK-Regular.ttc", "Noto Sans CJK"),
            ("NotoSansCJKsc-Regular.ttc", "Noto Sans CJK SC"),
            ("NotoSansCJKsc-Regular.otf", "Noto Sans CJK SC"),
            ("NotoSansSC-Regular.otf", "Noto Sans SC"),
            ("NotoSansSC-Regular.ttf", "Noto Sans SC"),
            # Windows
            ("msyh.ttc", "Microsoft YaHei"),
            ("msyhbd.ttc", "Microsoft YaHei Bold"),
            ("simhei.ttf", "SimHei"),
            ("simsun.ttc", "SimSun"),
            # macOS
            ("PingFang.ttc", "PingFang SC"),
            ("STHeiti Light.ttc", "STHeiti"),
            # Linux fallbacks
            ("DroidSansFallbackFull.ttf", "Droid Sans Fallback"),
            ("DejaVuSans.ttf", "DejaVu Sans"),
        ]

        for font_file, font_name in font_candidates:
            path = self._get_system_font_path(font_file)
            if path:
                self._unicode_font = font_name
                self._unicode_font_path = path
                break

    def _scan_for_cjk_font(self) -> str | None:
        """Scan font directories for any CJK-capable font."""
        font_dirs = [
            # Linux - where fonts-noto-cjk installs fonts
            "/usr/share/fonts/opentype/noto",
            "/usr/share/fonts/truetype/noto",
            "/usr/share/fonts/noto-cjk",
            "/usr/share/fonts/opentype",
            "/usr/share/fonts/truetype",
            "/usr/share/fonts",
        ]

        # Patterns that indicate CJK support
        cjk_patterns = ["notosanscjk", "notoserifcjk", "notosanssc", "notosc", "cjk"]

        for font_dir in font_dirs:
            if not os.path.exists(font_dir):
                continue
            for root, _dirs, files in os.walk(font_dir):
                for filename in files:
                    lower_name = filename.lower()
                    # Check if it's a font file with CJK support
                    if lower_name.endswith((".ttc", ".ttf", ".otf")):
                        for pattern in cjk_patterns:
                            if pattern in lower_name:
                                return os.path.join(root, filename)
        return None

    def _get_system_font_path(self, font_name: str) -> str | None:
        """Get the path to a system font file."""
        font_dirs = []

        if os.name == "nt":
            font_dirs.append(os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts"))

        font_dirs.extend(
            [
                # macOS
                "/System/Library/Fonts",
                "/Library/Fonts",
                os.path.expanduser("~/Library/Fonts"),
                # Linux - common font directories
                "/usr/share/fonts/opentype/noto",
                "/usr/share/fonts/truetype/noto",
                "/usr/share/fonts/noto-cjk",
                "/usr/share/fonts/opentype",
                "/usr/share/fonts/truetype",
                "/usr/share/fonts",
                "/usr/local/share/fonts",
                os.path.expanduser("~/.fonts"),
            ]
        )

        font_name_lower = font_name.lower()

        for font_dir in font_dirs:
            if os.path.exists(font_dir):
                # Direct path check
                font_path = os.path.join(font_dir, font_name)
                if os.path.exists(font_path):
                    return font_path
                # Walk and check with case-insensitive matching
                for root, _dirs, files in os.walk(font_dir):
                    for filename in files:
                        if filename.lower() == font_name_lower:
                            return os.path.join(root, filename)
        return None

    def render(
        self, nodes: list[DocumentNode], metadata: ExportMetadata | None = None
    ) -> bytes:
        """Render document nodes to PDF bytes."""
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Add Unicode font if available
        if self._unicode_font_path:
            pdf.add_font("unicode", style="", fname=self._unicode_font_path)
            pdf.add_font("unicode", style="B", fname=self._unicode_font_path)
            pdf.add_font("unicode", style="I", fname=self._unicode_font_path)
            font = "unicode"
        else:
            font = "Helvetica"

        pdf.add_page()
        pdf.set_font(font, size=11)

        # Render title page with cover, title, and metadata
        if metadata:
            has_cover = self._render_title_page(pdf, metadata, font)
            if has_cover:
                # Cover page: start content on a new page
                pdf.add_page()
            else:
                # No cover: add spacing after title/metadata before content
                pdf.ln(8)

        for node in nodes:
            self._render_node(pdf, node, font)

        return pdf.output()

    @staticmethod
    def _is_css_background(value: str) -> bool:
        """Check if a cover value is a CSS background (gradient or hex color)."""
        return value.startswith(("linear-gradient", "radial-gradient", "conic-gradient")) or bool(
            re.match(r"^#[0-9a-fA-F]{3,8}$", value)
        )

    @staticmethod
    def _parse_hex_color(hex_color: str) -> tuple[int, int, int]:
        """Parse a hex color string to (r, g, b) tuple."""
        h = hex_color.lstrip("#")
        if len(h) == 3:
            h = h[0] * 2 + h[1] * 2 + h[2] * 2
        elif len(h) < 6:
            h = h.ljust(6, "0")
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)

    def _parse_gradient_colors(self, gradient: str) -> list[tuple[int, int, int]]:
        """Extract color values from a CSS gradient string."""
        colors = []
        # Match hex colors
        for match in re.finditer(r"#[0-9a-fA-F]{3,8}", gradient):
            colors.append(self._parse_hex_color(match.group()))
        # Match rgb/rgba colors
        for match in re.finditer(r"rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)", gradient):
            colors.append((int(match.group(1)), int(match.group(2)), int(match.group(3))))
        return colors if colors else [(200, 200, 200)]

    @staticmethod
    def _download_image(url: str) -> str | None:
        """Download an image URL to a temporary file. Returns path or None."""
        import ipaddress
        from urllib.parse import urlparse

        def _is_private_host(host: str) -> bool:
            if not host:
                return True
            host_lower = host.lower()
            if host_lower in ("localhost", "ip6-localhost", "ip6-loopback"):
                return True
            try:
                ip = ipaddress.ip_address(host_lower)
                return (
                    ip.is_private
                    or ip.is_loopback
                    or ip.is_link_local
                    or ip.is_reserved
                    or ip.is_multicast
                )
            except ValueError:
                return False

        # Validate scheme
        parsed = urlparse(url)
        if parsed.scheme not in ("http", "https"):
            logger.warning("Blocked image download with non-HTTP scheme: %s", parsed.scheme)
            return None

        # SSRF protection: block private/internal hosts
        hostname = parsed.hostname or ""
        if _is_private_host(hostname):
            logger.warning("Blocked image download to private host: %s", hostname)
            return None

        try:
            import httpx

            resp = httpx.get(url, timeout=10, follow_redirects=True, max_redirects=5)
            if resp.status_code != 200:
                return None

            # Post-redirect SSRF check
            final_host = resp.url.host or ""
            if _is_private_host(final_host):
                logger.warning("Blocked image download after redirect to private host: %s", final_host)
                return None
            # Determine file extension from content type
            content_type = resp.headers.get("content-type", "")
            ext = ".jpg"
            if "png" in content_type:
                ext = ".png"
            elif "gif" in content_type:
                ext = ".gif"
            elif "webp" in content_type:
                ext = ".webp"
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                tmp.write(resp.content)
            return tmp.name
        except Exception as e:
            logger.warning(f"Failed to download cover image: {e}")
            return None

    def _render_cover_background(self, pdf: FPDF, cover_url: str, cover_height: float):
        """Render cover image or CSS background at the top of the page."""
        page_width = pdf.w
        x_start = 0
        y_start = 0

        if self._is_css_background(cover_url):
            if cover_url.startswith("#"):
                # Solid color
                r, g, b = self._parse_hex_color(cover_url)
                pdf.set_fill_color(r, g, b)
                pdf.rect(x_start, y_start, page_width, cover_height, "F")
            else:
                # CSS gradient — render as interpolated horizontal bands
                colors = self._parse_gradient_colors(cover_url)
                if len(colors) < 2:
                    colors = [colors[0], colors[0]]
                band_count = int(cover_height * 2)  # 2 bands per mm for smoothness
                for i in range(band_count):
                    t = i / max(band_count - 1, 1)
                    # Interpolate between first and last color
                    c0, c1 = colors[0], colors[-1]
                    r = int(c0[0] + (c1[0] - c0[0]) * t)
                    g = int(c0[1] + (c1[1] - c0[1]) * t)
                    b = int(c0[2] + (c1[2] - c0[2]) * t)
                    pdf.set_fill_color(r, g, b)
                    band_y = y_start + (i * cover_height / band_count)
                    band_h = cover_height / band_count + 0.1  # tiny overlap to avoid gaps
                    pdf.rect(x_start, band_y, page_width, band_h, "F")
        else:
            # Image URL — download and embed
            img_path = self._download_image(cover_url)
            if img_path:
                try:
                    pdf.image(img_path, x=x_start, y=y_start, w=page_width, h=cover_height)
                except Exception as e:
                    logger.warning(f"Failed to embed cover image in PDF: {e}")
                finally:
                    with contextlib.suppress(OSError):
                        os.unlink(img_path)

    def _render_title_page(self, pdf: FPDF, metadata: ExportMetadata, font: str) -> bool:
        """Render the title page with cover, title, and metadata.

        Returns True if a cover image was rendered (caller should add a page break).
        """
        has_cover = bool(metadata.cover_image_url)
        cover_height = 100.0  # mm

        if has_cover:
            self._render_cover_background(pdf, metadata.cover_image_url, cover_height)
            pdf.set_y(cover_height + 10)
        else:
            pdf.set_y(pdf.t_margin)

        # Title with optional icon
        title_text = metadata.title
        if metadata.icon:
            title_text = f"{metadata.icon}  {title_text}"

        pdf.set_font(font, "B", 24)
        pdf.set_text_color(26, 26, 26)
        pdf.multi_cell(0, 12, title_text)
        pdf.ln(4)

        # Separator line
        pdf.set_draw_color(200, 200, 200)
        y = pdf.get_y()
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(6)

        # Metadata block
        pdf.set_font(font, size=10)
        pdf.set_text_color(120, 120, 120)

        if metadata.author:
            pdf.cell(0, 5, f"Author: {metadata.author}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

        if metadata.created_at:
            date_str = metadata.created_at.strftime("%Y-%m-%d")
            pdf.cell(0, 5, f"Created: {date_str}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

        if metadata.updated_at:
            date_str = metadata.updated_at.strftime("%Y-%m-%d")
            pdf.cell(0, 5, f"Last updated: {date_str}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

        # Reset text color
        pdf.set_text_color(0)
        pdf.set_font(font, size=11)

        return has_cover

    def _render_node(self, pdf: FPDF, node: DocumentNode, font: str, level: int = 0):
        """Render a single node to the PDF."""
        if node.node_type == NodeType.TEXT:
            pdf.set_font(font, size=11)
            pdf.multi_cell(0, 6, node.content)
            pdf.ln(2)

        elif node.node_type == NodeType.HEADING:
            sizes = {1: 20, 2: 16, 3: 14, 4: 12, 5: 12, 6: 12}
            size = sizes.get(node.level, 12)
            pdf.set_font(font, "B", size)
            pdf.set_text_color(26, 26, 26)
            pdf.multi_cell(0, size / 2, node.content)
            pdf.ln(4 if node.level <= 2 else 2)
            pdf.set_text_color(0)

        elif node.node_type == NodeType.PARAGRAPH:
            pdf.set_font(font, size=11)
            text = self._get_paragraph_text(node)
            pdf.multi_cell(0, 6, text)
            pdf.ln(3)

        elif node.node_type == NodeType.CODE_BLOCK:
            pdf.set_font(font, size=9)
            pdf.set_fill_color(245, 245, 245)
            pdf.set_x(pdf.l_margin + 5)
            pdf.multi_cell(0, 5, node.content, fill=True)
            pdf.set_font(font, size=11)
            pdf.ln(3)

        elif node.node_type == NodeType.MERMAID_CHART:
            # Render mermaid source as code block fallback
            pdf.set_font(font, "B", 10)
            pdf.cell(0, 6, "[Mermaid Diagram]", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font(font, size=9)
            pdf.set_fill_color(245, 245, 245)
            pdf.set_x(pdf.l_margin + 5)
            pdf.multi_cell(0, 5, node.content, fill=True)
            pdf.set_font(font, size=11)
            pdf.ln(3)

        elif node.node_type == NodeType.BLOCKQUOTE:
            pdf.set_font(font, "I", 11)
            pdf.set_text_color(102, 102, 102)
            pdf.set_x(pdf.l_margin + 10)
            pdf.multi_cell(0, 6, node.content)
            pdf.set_text_color(0)
            pdf.ln(3)

        elif node.node_type == NodeType.LIST:
            self._render_list(pdf, node, font, level)
            pdf.ln(2)

        elif node.node_type == NodeType.TABLE:
            self._render_table(pdf, node, font)
            pdf.ln(3)

        elif node.node_type == NodeType.HORIZONTAL_RULE:
            pdf.set_draw_color(200, 200, 200)
            y = pdf.get_y()
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(5)

    def _get_paragraph_text(self, node: DocumentNode) -> str:
        """Extract text from paragraph with inline elements."""
        parts = []
        for child in node.children:
            parts.append(child.content)
        return "".join(parts) if parts else node.content

    def _render_list(self, pdf: FPDF, node: DocumentNode, font: str, level: int = 0):
        """Render a list to the PDF."""
        pdf.set_font(font, size=11)
        indent = 10 + (level * 10)
        counter = 1

        for item in node.children:
            if item.node_type == NodeType.LIST_ITEM:
                prefix = f"{counter}. " if node.ordered else "- "
                counter += 1
                pdf.set_x(pdf.l_margin + indent)
                pdf.multi_cell(0, 6, prefix + item.content)

                # Handle nested lists
                for nested in item.children:
                    if nested.node_type == NodeType.LIST:
                        self._render_list(pdf, nested, font, level + 1)

    def _render_table(self, pdf: FPDF, node: DocumentNode, font: str):
        """Render a table to the PDF with auto-wrapping and proportional column widths."""
        if not node.children:
            return

        first_row = node.children[0]
        col_count = len(first_row.children)
        if col_count == 0:
            return

        available_width = pdf.w - pdf.l_margin - pdf.r_margin
        line_height = 5
        padding = 1.5
        font_size = 9

        pdf.set_font(font, size=font_size)

        # Calculate proportional column widths based on max text width per column
        col_max_text_w = [0.0] * col_count
        for row in node.children:
            for j, cell in enumerate(row.children):
                if j < col_count:
                    w = pdf.get_string_width(cell.content or "")
                    col_max_text_w[j] = max(col_max_text_w[j], w)

        min_col_w = 15.0
        # Cap each column's contribution so one very long cell can't starve others
        max_col_w = max(available_width / 2, min_col_w)
        raw_widths = [max(min(w + padding * 4, max_col_w), min_col_w) for w in col_max_text_w]
        total_raw = sum(raw_widths)
        col_widths = [w * available_width / total_raw for w in raw_widths]
        # Enforce hard minimum to prevent fpdf2 "not enough horizontal space" error
        hard_min = padding * 2 + 2  # must be > padding*2 for text to fit
        col_widths = [max(w, hard_min) for w in col_widths]
        total_w = sum(col_widths)
        if total_w > available_width:
            col_widths = [max(w * available_width / total_w, hard_min) for w in col_widths]

        pdf.set_draw_color(180, 180, 180)

        for i, row in enumerate(node.children):
            is_header_row = i == 0

            # Collect cell texts for this row
            cells_text: list[str] = []
            cells_is_header: list[bool] = []
            for j in range(col_count):
                if j < len(row.children):
                    cells_text.append(row.children[j].content or "")
                    cells_is_header.append(row.children[j].is_header or is_header_row)
                else:
                    cells_text.append("")
                    cells_is_header.append(is_header_row)

            # First pass: calculate max row height
            max_cell_h = line_height + padding * 2
            for j, text in enumerate(cells_text):
                pdf.set_font(font, "B" if cells_is_header[j] else "", font_size)
                usable_w = col_widths[j] - padding * 2
                if usable_w <= 0:
                    usable_w = 1
                text_w = pdf.get_string_width(text) if text else 0
                num_lines = max(1, math.ceil(text_w / usable_w)) if text_w > 0 else 1
                cell_h = num_lines * line_height + padding * 2
                max_cell_h = max(max_cell_h, cell_h)

            # Page break check
            if pdf.get_y() + max_cell_h > pdf.h - pdf.b_margin:
                pdf.add_page()

            # Second pass: render cells
            y_start = pdf.get_y()
            x = pdf.l_margin

            for j, text in enumerate(cells_text):
                w = col_widths[j]
                is_header = cells_is_header[j]

                if is_header:
                    pdf.set_font(font, "B", font_size)
                    pdf.set_fill_color(240, 240, 240)
                else:
                    pdf.set_font(font, "", font_size)
                    pdf.set_fill_color(255, 255, 255)

                # Draw cell background + border
                pdf.rect(x, y_start, w, max_cell_h, "DF")

                # Render wrapped text inside cell
                pdf.set_xy(x + padding, y_start + padding)
                pdf.multi_cell(max(w - padding * 2, 1.0), line_height, text)

                x += w

            pdf.set_y(y_start + max_cell_h)


# ============================================================================
# DOCX Renderer
# ============================================================================


class DOCXRenderer:
    """Renders DocumentNodes to DOCX."""

    def render(
        self, nodes: list[DocumentNode], metadata: ExportMetadata | None = None
    ) -> bytes:
        """Render document nodes to DOCX bytes."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Render title page if metadata is provided
        if metadata:
            self._render_title_page(doc, metadata)

        for node in nodes:
            self._render_node(doc, node)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _render_title_page(self, doc: Document, metadata: ExportMetadata):
        """Render title page with cover image, title, and metadata."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        # Cover image (only for URL-based images, skip CSS gradients)
        if metadata.cover_image_url and not metadata.cover_image_url.startswith(
            ("linear-gradient", "radial-gradient", "conic-gradient", "#")
        ):
            img_path = PDFRenderer._download_image(metadata.cover_image_url)
            if img_path:
                try:
                    doc.add_picture(img_path, width=Inches(6.5))
                except Exception as e:
                    logger.warning(f"Failed to embed cover image in DOCX: {e}")
                finally:
                    with contextlib.suppress(OSError):
                        os.unlink(img_path)

        # Title
        title_text = metadata.title
        if metadata.icon:
            title_text = f"{metadata.icon}  {title_text}"
        title_para = doc.add_heading(title_text, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Separator
        sep = doc.add_paragraph()
        sep.add_run("_" * 60)
        sep.runs[0].font.color.rgb = RGBColor(200, 200, 200)
        sep.runs[0].font.size = Pt(8)

        # Metadata
        if metadata.author:
            p = doc.add_paragraph()
            run = p.add_run(f"Author: {metadata.author}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(120, 120, 120)

        if metadata.created_at:
            p = doc.add_paragraph()
            run = p.add_run(f"Created: {metadata.created_at.strftime('%Y-%m-%d')}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(120, 120, 120)

        if metadata.updated_at:
            p = doc.add_paragraph()
            run = p.add_run(f"Last updated: {metadata.updated_at.strftime('%Y-%m-%d')}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(120, 120, 120)

        # Page break after title page (only if there's a cover or metadata)
        if metadata.cover_image_url:
            doc.add_page_break()

    def _render_node(self, doc: Document, node: DocumentNode, level: int = 0):
        """Render a single node to the document."""
        from docx.shared import Inches, Pt, RGBColor

        if node.node_type == NodeType.TEXT:
            doc.add_paragraph(node.content)

        elif node.node_type == NodeType.HEADING:
            doc.add_heading(node.content, level=node.level)

        elif node.node_type == NodeType.PARAGRAPH:
            para = doc.add_paragraph()
            self._render_inline_content(para, node.children)

        elif node.node_type == NodeType.CODE_BLOCK:
            para = doc.add_paragraph()
            run = para.add_run(node.content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            para.paragraph_format.left_indent = Inches(0.25)

        elif node.node_type == NodeType.MERMAID_CHART:
            # Render mermaid source as code block fallback
            para = doc.add_paragraph()
            run = para.add_run("[Mermaid Diagram]")
            run.bold = True
            para = doc.add_paragraph()
            run = para.add_run(node.content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            para.paragraph_format.left_indent = Inches(0.25)

        elif node.node_type == NodeType.BLOCKQUOTE:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run(node.content)
            run.font.italic = True
            run.font.color.rgb = RGBColor(102, 102, 102)

        elif node.node_type == NodeType.LIST:
            self._render_list(doc, node, level)

        elif node.node_type == NodeType.TABLE:
            self._render_table(doc, node)

        elif node.node_type == NodeType.HORIZONTAL_RULE:
            para = doc.add_paragraph()
            para.add_run("_" * 50)

    def _render_inline_content(self, para, children: list[DocumentNode]):
        """Render inline content to a paragraph."""
        from docx.shared import Pt, RGBColor

        for child in children:
            if child.node_type == NodeType.TEXT:
                para.add_run(child.content)
            elif child.node_type == NodeType.INLINE_BOLD:
                run = para.add_run(child.content)
                run.bold = True
            elif child.node_type == NodeType.INLINE_ITALIC:
                run = para.add_run(child.content)
                run.italic = True
            elif child.node_type == NodeType.INLINE_CODE:
                run = para.add_run(child.content)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            elif child.node_type == NodeType.INLINE_LINK:
                run = para.add_run(child.content)
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.underline = True

    def _render_list(self, doc: Document, node: DocumentNode, level: int = 0):
        """Render a list to the document."""
        from docx.shared import Inches

        counter = 1
        for item in node.children:
            if item.node_type == NodeType.LIST_ITEM:
                prefix = f"{counter}. " if node.ordered else "- "
                counter += 1
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                para.add_run(prefix + item.content)

                for nested in item.children:
                    if nested.node_type == NodeType.LIST:
                        self._render_list(doc, nested, level + 1)

    def _render_table(self, doc: Document, node: DocumentNode):
        """Render a table to the document."""
        if not node.children:
            return

        col_count = len(node.children[0].children)
        if col_count == 0:
            return

        table = doc.add_table(rows=len(node.children), cols=col_count)
        table.style = "Table Grid"

        for i, row in enumerate(node.children):
            for j, cell in enumerate(row.children):
                if j < col_count:
                    table_cell = table.rows[i].cells[j]
                    table_cell.text = cell.content

                    if cell.is_header:
                        for para in table_cell.paragraphs:
                            for run in para.runs:
                                run.bold = True


# ============================================================================
# Export Service (Public Interface)
# ============================================================================


class ExportService:
    """Service for exporting content to various formats."""

    def __init__(self):
        import markdown

        self.md = markdown.Markdown(extensions=["tables", "fenced_code", "toc"])
        self.parser = HTMLToDocumentParser()
        self.pdf_renderer = PDFRenderer()
        self.docx_renderer = DOCXRenderer()

    def export_markdown(self, content: str, filename: str) -> bytes:  # noqa: ARG002
        """Export content as Markdown."""
        from markdownify import markdownify as md

        markdown_content = md(
            content, heading_style="ATX", code_language_callback=self._get_code_language
        )
        return markdown_content.encode("utf-8")

    @staticmethod
    def _get_code_language(el):
        """Extract language from code block class attribute (e.g. 'language-python')."""
        classes = el.get("class") or []
        for cls in classes:
            if cls.startswith("language-"):
                return cls[len("language-") :]
        return None

    def export_pdf(
        self, content: str, filename: str, metadata: dict | None = None  # noqa: ARG002
    ) -> bytes:
        """Export content as PDF."""
        self.md.reset()
        html = self.md.convert(content)
        nodes = self.parser.parse(html)
        meta = ExportMetadata(**metadata) if metadata else None
        return self.pdf_renderer.render(nodes, metadata=meta)

    def export_docx(
        self, content: str, filename: str, metadata: dict | None = None  # noqa: ARG002
    ) -> bytes:
        """Export content as DOCX."""
        self.md.reset()
        html = self.md.convert(content)
        nodes = self.parser.parse(html)
        meta = ExportMetadata(**metadata) if metadata else None
        return self.docx_renderer.render(nodes, metadata=meta)


# Lazy singleton — instantiated on first use, not at import time
_export_service: ExportService | None = None


def get_export_service() -> ExportService:
    """Get cached export service instance."""
    global _export_service
    if _export_service is None:
        _export_service = ExportService()
    return _export_service
