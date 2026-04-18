"""File converter service.

Converts PDF, DOCX, PPTX files to Markdown using the active LLM provider's
vision-capable model (mapped via the ``file_conversion`` role). Falls
back to markitdown when no provider is configured or the model isn't
vision-capable.
"""

import asyncio
import base64
import io
import logging

from openai import AsyncOpenAI

from config import get_settings
from provider.registry import (
    ProviderUnconfiguredError,
    active_provider_id,
    build_client,
    provider_api_key,
    role_model,
)

logger = logging.getLogger(__name__)

# Supported file types for conversion
SUPPORTED_TYPES = {".pdf", ".docx", ".pptx"}

# Types that need pre-conversion to PDF via LibreOffice before LLM processing
OFFICE_TYPES = {".docx", ".pptx"}

# Conversion prompt for native file uploads (PDF, images)
CONVERSION_PROMPT = """Convert this document to well-formatted Markdown.

Requirements:
- Preserve all text content accurately
- Maintain the document structure (headings, lists, tables, etc.)
- Use appropriate Markdown syntax for formatting
- For tables, use proper Markdown table syntax
- For code blocks, use fenced code blocks with language hints if identifiable
- Preserve any mathematical formulas using LaTeX syntax ($...$ for inline, $$...$$ for block)
- Do not add any commentary or explanations, just output the converted Markdown
- Do NOT wrap the output in code fences (```). Output raw Markdown directly
- If the document contains images with text, include the text content

Output only the Markdown content, nothing else."""

# Conversion prompt for pre-extracted text (DOCX, PPTX)
TEXT_CONVERSION_PROMPT = """Format the following extracted document content into well-structured Markdown.

Requirements:
- Preserve all text content accurately
- Infer and apply appropriate heading levels based on context
- Use appropriate Markdown syntax for formatting (bold, italic, lists, etc.)
- For tables, use proper Markdown table syntax
- For code blocks, use fenced code blocks with language hints if identifiable
- Preserve any mathematical formulas using LaTeX syntax ($...$ for inline, $$...$$ for block)
- Do not add any commentary or explanations, just output the formatted Markdown
- Do NOT wrap the output in code fences (```). Output raw Markdown directly
- Clean up any formatting artifacts from the extraction process

Document content:
"""

# Conversion prompt for chunked PDF pages (parallel processing)
CHUNK_CONVERSION_PROMPT = """Convert these pages to well-formatted Markdown.

Requirements:
- Preserve all text content accurately
- Maintain document structure (headings, lists, tables)
- Use proper Markdown syntax
- Preserve mathematical formulas using LaTeX syntax ($...$ for inline, $$...$$ for block)
- Do NOT wrap output in code fences (```). Output raw Markdown directly
- If content continues from a previous page, continue naturally without repeating headings
- Output only the Markdown content"""

# Parallel PDF conversion settings
PARALLEL_PDF_PAGE_THRESHOLD = 5  # Only split PDFs with more than this many pages
PAGES_PER_CHUNK = 5  # Pages per parallel chunk
MAX_CONCURRENT_CHUNKS = 10  # Max concurrent API calls


def _build_extra_body() -> dict | None:
    """Hook for provider-specific extra_body. No-op today."""
    return None


def _resolve_file_conversion() -> tuple[AsyncOpenAI, str]:
    """Return (client, model) for the file_conversion role on the active provider."""
    pid = active_provider_id()
    if pid is None:
        raise ProviderUnconfiguredError(
            "No LLM provider configured. Open Settings to add an API key."
        )
    key = provider_api_key(pid) or get_settings().env_api_key_for(pid)
    if not key:
        raise ProviderUnconfiguredError(f"Active provider '{pid}' has no API key.")
    model = role_model("file_conversion", pid)
    if not model:
        raise ProviderUnconfiguredError(
            f"Provider '{pid}' has no file_conversion model configured."
        )
    return build_client(key, pid), model


def _get_client(api_key: str | None = None) -> AsyncOpenAI:
    """Build a client for the active provider (explicit key override allowed)."""
    pid = active_provider_id()
    if pid is None:
        raise ProviderUnconfiguredError(
            "No LLM provider configured. Open Settings to add an API key."
        )
    effective_key = api_key or provider_api_key(pid) or get_settings().env_api_key_for(pid)
    if not effective_key:
        raise ProviderUnconfiguredError(f"Active provider '{pid}' has no API key.")
    return build_client(effective_key, pid)


def _find_libreoffice() -> str:
    """Find the LibreOffice executable path across platforms."""
    import shutil
    import sys

    # Try common command names (Linux/macOS: libreoffice or soffice, Windows: soffice)
    for cmd in ("libreoffice", "soffice"):
        path = shutil.which(cmd)
        if path:
            return path

    # Windows: check common install locations
    if sys.platform == "win32":
        import glob

        for pattern in [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]:
            matches = glob.glob(pattern)
            if matches:
                return matches[0]

    raise RuntimeError(
        "LibreOffice not found. Install it: "
        "apt-get install libreoffice-core (Linux) or https://www.libreoffice.org (Windows/macOS)"
    )


def _convert_office_to_pdf(content: bytes, extension: str) -> bytes:
    """Convert DOCX/PPTX to PDF using LibreOffice headless.

    Args:
        content: Raw file bytes
        extension: File extension (e.g., '.docx', '.pptx')

    Returns:
        PDF file bytes

    Raises:
        RuntimeError: If LibreOffice is not installed or conversion fails
    """
    import os
    import subprocess
    import tempfile

    soffice = _find_libreoffice()

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, f"input{extension}")
        with open(input_path, "wb") as f:
            f.write(content)

        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, input_path],
            capture_output=True,
            timeout=120,
        )

        if result.returncode != 0:
            stderr = result.stderr.decode(errors="replace")
            raise RuntimeError(f"LibreOffice conversion failed: {stderr}")

        pdf_path = os.path.join(tmpdir, "input.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("LibreOffice did not produce PDF output")

        with open(pdf_path, "rb") as f:
            return f.read()


def extract_docx_content(content: bytes) -> str:
    """Extract text content from a DOCX file.

    Args:
        content: Raw DOCX file bytes

    Returns:
        Extracted text content with basic structure preserved
    """
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Try to preserve heading structure based on style
        style_name = para.style.name.lower() if para.style else ""
        if "heading" in style_name:
            # Extract heading level from style name (e.g., "Heading 1" -> 1)
            try:
                level = int("".join(filter(str.isdigit, style_name)) or "1")
                level = min(level, 6)  # Cap at h6
                parts.append(f"{'#' * level} {text}")
            except ValueError:
                parts.append(f"# {text}")
        elif "title" in style_name:
            parts.append(f"# {text}")
        else:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append("| " + " | ".join(cells) + " |")

        if table_rows:
            # Add header separator after first row
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            table_rows.insert(1, header_sep)
            parts.append("\n".join(table_rows))

    return "\n\n".join(parts)


async def convert_file_to_markdown(
    content: bytes,
    filename: str,
    extension: str,
    model: str | None = None,
    api_key: str | None = None,
) -> tuple[str, dict | None]:
    """Convert file content to Markdown via OpenRouter API.

    Args:
        content: Raw file bytes
        filename: Original filename (for logging)
        extension: File extension (e.g., '.pdf', '.docx')
        model: Model to use for conversion (defaults to the file_conversion role)
        api_key: Optional user API key (falls back to active provider's key)

    Returns:
        Tuple of:
        - Converted Markdown content
        - Usage dict: {model, input_tokens, output_tokens, cost, is_byok}

    Raises:
        ValueError: If file type is not supported or API key not configured
        Exception: If conversion fails
    """
    ext_lower = extension.lower()

    if ext_lower not in SUPPORTED_TYPES:
        raise ValueError(f"Unsupported file type: {extension}")

    settings = get_settings()
    effective_model = model or role_model("file_conversion") or ""
    if not effective_model:
        raise ProviderUnconfiguredError(
            "No file_conversion model configured for the active provider."
        )

    # For DOCX/PPTX: convert to PDF first via LibreOffice, then process as PDF
    if ext_lower in OFFICE_TYPES:
        try:
            logger.info(f"Converting {filename} to PDF via LibreOffice")
            content = await asyncio.to_thread(_convert_office_to_pdf, content, ext_lower)
            ext_lower = ".pdf"
            logger.info(f"Successfully converted {filename} to PDF")
        except Exception as e:
            logger.warning(
                f"LibreOffice conversion failed for {filename}: {e}. Falling back to markitdown."
            )
            fallback_text = await markitdown_convert(content, filename, extension.lower())
            return fallback_text, None

    # From here on, content is always PDF
    mime_type = "application/pdf"

    # For large PDFs, use parallel chunk conversion for speed
    num_pages = _get_pdf_page_count(content)
    if num_pages > PARALLEL_PDF_PAGE_THRESHOLD:
        try:
            logger.info(f"{filename} has {num_pages} pages, using parallel conversion")
            return await _convert_pdf_parallel(content, filename, effective_model, api_key)
        except Exception as e:
            logger.warning(
                f"Parallel PDF conversion failed for {filename}, trying single-call: {e}"
            )
            # Fall through to single-call below

    # Single-call conversion for small PDFs or parallel fallback
    try:
        client = _get_client(api_key)
        b64_data = base64.b64encode(content).decode("utf-8")
        data_url = f"data:{mime_type};base64,{b64_data}"

        response = await client.chat.completions.create(
            model=effective_model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": data_url},
                        },
                        {
                            "type": "text",
                            "text": CONVERSION_PROMPT,
                        },
                    ],
                }
            ],
            max_tokens=settings.file_conversion_max_tokens,
            extra_body=_build_extra_body(),
        )

        from services.usage_tracker import extract_usage

        usage = extract_usage(response)
        usage["model"] = effective_model
        usage["is_byok"] = bool(api_key)

        markdown_content = response.choices[0].message.content

        # Detect truncation: finish_reason == "length" means output hit max_tokens
        finish_reason = getattr(response.choices[0], "finish_reason", None)
        if finish_reason == "length":
            logger.warning(
                f"LLM output truncated for {filename} (finish_reason=length). "
                f"Falling back to markitdown."
            )
            fallback_text = await markitdown_convert(content, filename, ext_lower)
            return fallback_text, None

        if not markdown_content:
            raise ValueError("LLM returned empty response")

        logger.info(f"Successfully converted {filename} to Markdown")
        return markdown_content, usage

    except Exception as e:
        logger.warning(f"LLM conversion failed for {filename}, falling back to markitdown: {e}")
        fallback_text = await markitdown_convert(content, filename, ext_lower)
        return fallback_text, None


async def _convert_docx_to_markdown(
    content: bytes, filename: str, model: str, api_key: str | None = None
) -> tuple[str, dict | None]:
    """Convert DOCX file to Markdown by extracting text and formatting with LLM.

    Args:
        content: Raw DOCX file bytes
        filename: Original filename (for logging)
        model: Model to use for formatting
        api_key: Optional user API key

    Returns:
        Formatted Markdown content
    """
    # Extract text content from DOCX via markitdown (much more complete than python-docx)
    try:
        extracted_text = await markitdown_convert(content, filename, ".docx")
    except Exception as e:
        logger.warning(f"markitdown failed for {filename}, trying python-docx: {e}")
        try:
            extracted_text = await asyncio.to_thread(extract_docx_content, content)
        except Exception as e2:
            logger.error(f"Failed to extract text from {filename}: {e2}")
            raise ValueError(f"Failed to read DOCX file: {e2}")

    if not extracted_text.strip():
        raise ValueError("DOCX file appears to be empty")

    # Use LLM to format the extracted text as proper Markdown
    try:
        settings = get_settings()
        client = _get_client(api_key)
        response = await client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": TEXT_CONVERSION_PROMPT + extracted_text}],
            max_tokens=settings.file_conversion_max_tokens,
            extra_body=_build_extra_body(),
        )

        from services.usage_tracker import extract_usage

        usage = extract_usage(response)
        usage["model"] = model
        usage["is_byok"] = bool(api_key)

        markdown_content = response.choices[0].message.content

        # Detect truncation: finish_reason == "length" means output hit max_tokens
        finish_reason = getattr(response.choices[0], "finish_reason", None)
        if finish_reason == "length":
            logger.warning(
                f"LLM output truncated for {filename} (finish_reason=length). "
                f"Falling back to markitdown."
            )
            fallback_text = await markitdown_convert(content, filename, ".docx")
            return fallback_text, None

        if not markdown_content:
            raise ValueError("LLM returned empty response")

        logger.info(f"Successfully converted {filename} to Markdown")
        return markdown_content, usage

    except Exception as e:
        logger.warning(f"LLM formatting failed for {filename}, falling back to markitdown: {e}")
        fallback_text = await markitdown_convert(content, filename, ".docx")
        return fallback_text, None


def _split_pdf_to_chunks(content: bytes, pages_per_chunk: int = PAGES_PER_CHUNK) -> list[bytes]:
    """Split PDF into chunks of N pages, each returned as separate PDF bytes.

    Args:
        content: Raw PDF file bytes
        pages_per_chunk: Number of pages per chunk

    Returns:
        List of PDF byte chunks
    """
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(content))
    chunks = []
    for start in range(0, len(reader.pages), pages_per_chunk):
        writer = PdfWriter()
        for page in reader.pages[start : start + pages_per_chunk]:
            writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        chunks.append(buf.getvalue())
    return chunks


def _get_pdf_page_count(content: bytes) -> int:
    """Get the number of pages in a PDF. Returns 1 if page count cannot be determined."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        return len(reader.pages)
    except Exception:
        return 1


async def _convert_pdf_parallel(
    content: bytes, filename: str, model: str, api_key: str | None = None
) -> tuple[str, dict | None]:
    """Convert large PDF by splitting into chunks and converting in parallel.

    Splits the PDF into chunks of PAGES_PER_CHUNK pages, sends each chunk
    to the LLM concurrently (up to MAX_CONCURRENT_CHUNKS at a time),
    then merges the results in order.

    Args:
        content: Raw PDF file bytes
        filename: Original filename (for logging)
        model: Model to use for conversion
        api_key: Optional user API key

    Returns:
        Tuple of merged markdown content and aggregated usage dict
    """
    settings = get_settings()
    chunks = await asyncio.to_thread(_split_pdf_to_chunks, content)
    num_chunks = len(chunks)

    logger.info(f"Splitting {filename} into {num_chunks} chunks for parallel conversion")

    semaphore = asyncio.Semaphore(MAX_CONCURRENT_CHUNKS)
    total_usage: dict = {"input_tokens": 0, "output_tokens": 0, "cost": 0.0}

    # Per-chunk token limit: proportional to total, with 2x headroom
    chunk_max_tokens = max(4096, settings.file_conversion_max_tokens // num_chunks * 2)

    async def convert_chunk(idx: int, chunk_bytes: bytes) -> tuple[int, str]:
        async with semaphore:
            client = _get_client(api_key)
            b64 = base64.b64encode(chunk_bytes).decode("utf-8")
            data_url = f"data:application/pdf;base64,{b64}"

            response = await client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "image_url", "image_url": {"url": data_url}},
                            {"type": "text", "text": CHUNK_CONVERSION_PROMPT},
                        ],
                    }
                ],
                max_tokens=chunk_max_tokens,
                extra_body=_build_extra_body(),
            )

            # Accumulate usage
            from services.usage_tracker import extract_usage

            chunk_usage = extract_usage(response)
            total_usage["input_tokens"] += chunk_usage.get("input_tokens") or 0
            total_usage["output_tokens"] += chunk_usage.get("output_tokens") or 0
            total_usage["cost"] += chunk_usage.get("cost") or 0.0

            chunk_text = response.choices[0].message.content or ""
            return idx, chunk_text

    # Run all chunks in parallel
    tasks = [convert_chunk(i, chunk) for i, chunk in enumerate(chunks)]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    # Separate successes from failures
    successful: list[tuple[int, str]] = []
    failed_count = 0
    for r in results:
        if isinstance(r, Exception):
            failed_count += 1
            logger.warning(f"Chunk conversion failed for {filename}: {r}")
        else:
            successful.append(r)

    if not successful:
        raise ValueError(f"All {num_chunks} chunk conversions failed for {filename}")

    if failed_count > 0:
        logger.warning(
            f"{failed_count}/{num_chunks} chunks failed for {filename}, "
            f"merging {len(successful)} successful chunks"
        )

    # Sort by chunk index and merge
    successful.sort(key=lambda x: x[0])
    merged = "\n\n".join(text.strip() for _, text in successful if text.strip())

    total_usage["model"] = model
    total_usage["is_byok"] = bool(api_key)

    logger.info(
        f"Parallel conversion of {filename} complete: "
        f"{len(successful)}/{num_chunks} chunks, "
        f"{total_usage['output_tokens']} output tokens"
    )

    return merged, total_usage


async def markitdown_convert(content: bytes, filename: str, ext: str) -> str:
    """Fallback conversion via markitdown (no LLM needed).

    Uses Microsoft's markitdown library to convert PDF, DOCX, PPTX to Markdown.
    Lower quality than LLM conversion but zero cost.
    """
    import os
    import tempfile

    from markitdown import MarkItDown

    md_converter = MarkItDown(enable_plugins=False)

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    try:
        result = await asyncio.to_thread(md_converter.convert, tmp_path)
        text = result.text_content
        if not text or not text.strip():
            raise ValueError(f"markitdown returned empty content for {filename}")
        logger.info(f"Successfully converted {filename} via markitdown fallback")
        return text
    finally:
        os.unlink(tmp_path)


def is_converter_configured() -> bool:
    """Check if the active provider has an API key and a file_conversion model."""
    pid = active_provider_id()
    if pid is None:
        return False
    key = provider_api_key(pid) or get_settings().env_api_key_for(pid)
    return bool(key and role_model("file_conversion", pid))
